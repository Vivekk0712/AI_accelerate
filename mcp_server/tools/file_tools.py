import os
import uuid
import hashlib
from typing import List, Dict, Optional, Any
from datetime import datetime
import PyPDF2
import io
import json
import logging

logger = logging.getLogger(__name__)

# Import enhanced embedding functions
try:
    from embeddings import generate_embedding, generate_embeddings_batch, rerank_results, EMBEDDING_DIM
    SEMANTIC_EMBEDDINGS_AVAILABLE = True
    print("✅ Semantic embeddings enabled (Sentence Transformers)")
except ImportError:
    print("⚠️  Semantic embeddings not available, using fallback hash-based embeddings")
    SEMANTIC_EMBEDDINGS_AVAILABLE = False
    EMBEDDING_DIM = 384  # Match Sentence Transformers dimension
    
    def generate_embedding(text: str) -> List[float]:
        """Fallback hash-based embedding if Sentence Transformers not available"""
        try:
            import hashlib
            embeddings = []
            base_text = text if text else "empty"
            
            for i in range(EMBEDDING_DIM):
                hash_input = f"{base_text}_{i}_{len(base_text)}".encode()
                if i % 4 == 0:
                    hash_obj = hashlib.md5(hash_input)
                elif i % 4 == 1:
                    hash_obj = hashlib.sha1(hash_input)
                elif i % 4 == 2:
                    hash_obj = hashlib.sha256(hash_input)
                else:
                    hash_obj = hashlib.blake2b(hash_input, digest_size=8)
                
                hash_bytes = hash_obj.digest()
                byte_value = hash_bytes[i % len(hash_bytes)]
                embeddings.append(float(byte_value) / 255.0)
            
            return embeddings
        except Exception as e:
            print(f"Error generating embedding: {e}")
            return [0.0] * EMBEDDING_DIM
    
    def generate_embeddings_batch(texts: List[str], batch_size: int = 32) -> List[List[float]]:
        """Fallback batch embedding"""
        return [generate_embedding(text) for text in texts]
    
    def rerank_results(query: str, documents: List[str], top_k: Optional[int] = None) -> List[tuple]:
        """Fallback re-ranking (no-op)"""
        return [(idx, 0.5) for idx in range(len(documents))]

def extract_text_from_file(file_content: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract text from various file types
    """
    import os
    import mimetypes

    # Detect MIME type from filename
    mime_type, _ = mimetypes.guess_type(filename)
    if not mime_type:
        mime_type = 'application/octet-stream'

    try:
        if mime_type == 'application/pdf':
            from PyPDF2 import PdfReader
            import io
            pdf_reader = PdfReader(io.BytesIO(file_content))
            text_content = "\n".join([page.extract_text() or "" for page in pdf_reader.pages])
            return {
                'text': text_content,
                'mime_type': mime_type,
                'chunks': _chunk_text(text_content),
                'page_count': len(pdf_reader.pages)
            }
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or filename.lower().endswith('.docx'):
            from docx import Document
            import io
            doc = Document(io.BytesIO(file_content))
            text_content = "\n".join([para.text for para in doc.paragraphs if para.text])
            return {
                'text': text_content,
                'mime_type': mime_type,
                'chunks': _chunk_text(text_content),
                'paragraph_count': len(doc.paragraphs)
            }
        elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' or filename.lower().endswith('.xlsx'):
            import openpyxl
            import io
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
            text_content = []
            for sheet in workbook:
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join(str(cell) for cell in row if cell is not None)
                    if row_text:
                        text_content.append(row_text)
            text_content = "\n".join(text_content)
            return {
                'text': text_content,
                'mime_type': mime_type,
                'chunks': _chunk_text(text_content),
                'sheet_count': len(workbook.sheetnames)
            }
        elif mime_type in ['text/plain', 'text/html', 'application/json', 'text/csv', 'application/xml', 'text/xml'] or any(filename.lower().endswith(ext) for ext in ['.txt','.html','.json','.csv','.xml']):
            text_content = file_content.decode('utf-8', errors='ignore')
            if mime_type == 'text/html' or filename.lower().endswith('.html'):
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(text_content, 'lxml')
                text_content = soup.get_text()
            return {
                'text': text_content,
                'mime_type': mime_type,
                'chunks': _chunk_text(text_content)
            }
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
    except Exception as e:
        raise Exception(f"Error extracting text from {filename}: {str(e)}")

def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200, default_page: int = 1) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks and return structured chunk dicts.
    """
    chunks: List[Dict[str, Any]] = []
    start = 0
    chunk_index = 0
    text = text or ""
    while start < len(text):
        end = start + chunk_size
        chunk_content = text[start:end]
        chunks.append({
            'chunk_index': chunk_index,
            'content': chunk_content,
            'page_number': default_page
        })
        start = max(end - overlap, end) if overlap > 0 else end
        chunk_index += 1
    # Handle empty text - still create a single empty chunk
    if not chunks:
        chunks.append({
            'chunk_index': 0,
            'content': '',
            'page_number': default_page
        })
    return chunks

def upload_file_to_storage(file_content: bytes, filename: str, user_id: str) -> str:
    """Upload file to Supabase Storage with MIME type detection"""
    try:
        from supabase_client import supabase # Local import
        if supabase is None:
            raise Exception("Supabase client not initialized")
        import mimetypes
        mime_type, _ = mimetypes.guess_type(filename)
        if not mime_type:
            mime_type = 'application/octet-stream'
        allowed_mime_types = [
            'application/pdf',
            'text/plain',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
            'text/html',
            'application/json',
            'text/csv',
            'application/xml',
            'text/xml'
        ]
        if mime_type not in allowed_mime_types:
            raise ValueError(f"Unsupported file type: {mime_type}")
        file_extension = os.path.splitext(filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        storage_path = f"uploads/{user_id}/{unique_filename}"
        response = supabase.storage.from_("files").upload(
            path=storage_path,
            file=file_content,
            file_options={"content-type": mime_type}
        )
        if hasattr(response, 'status_code') and response.status_code != 200:
            raise Exception(f"Storage upload failed with status {response.status_code}")
        return storage_path
    except Exception as e:
        raise Exception(f"Failed to upload file to storage: {str(e)}")

def create_file_record(user_id: str, filename: str, file_size: int, file_path: str, content_type: str) -> Dict[str, Any]:
    """Create file record in database"""
    try:
        from supabase_client import supabase
        if supabase is None:
            raise Exception("Supabase client not initialized")
        file_extension = os.path.splitext(filename)[1].lstrip('.').lower() or 'bin'
        file_data = {
            'user_id': user_id,
            'filename': os.path.basename(file_path),
            'original_filename': filename,
            'file_type': file_extension,
            'file_size': file_size,
            'file_path': file_path,
            'content_type': content_type,
            'upload_status': 'uploaded'
        }
        response = supabase.table('files').insert(file_data).execute()
        if response.data:
            return response.data[0]
        else:
            raise Exception("Failed to create file record")
    except Exception as e:
        raise Exception(f"Failed to create file record: {str(e)}")

def process_file_chunks(file_id: str, chunks: List[Dict[str, Any]], user_id: str, filename: str = None) -> List[Dict[str, Any]]:
    """Process and store file chunks with embeddings in Elasticsearch"""
    try:
        from supabase_client import supabase
        from elasticsearch_client import index_document_chunk
        
        if supabase is None:
            raise Exception("Supabase client not initialized")
        
        chunk_records: List[Dict[str, Any]] = []
        
        for chunk in chunks:
            # Ensure chunk is a dict
            if isinstance(chunk, str):
                chunk = {
                    'chunk_index': len(chunk_records),
                    'content': chunk,
                    'page_number': None
                }
            
            # Store chunk metadata in Supabase
            chunk_data = {
                'file_id': file_id,
                'chunk_index': chunk.get('chunk_index', len(chunk_records)),
                'content': chunk.get('content', ''),
                'page_number': chunk.get('page_number')
            }
            chunk_response = supabase.table('file_chunks').insert(chunk_data).execute()
            
            if not chunk_response.data:
                continue
            
            chunk_record = chunk_response.data[0]
            chunk_records.append(chunk_record)
            
            # Generate embedding
            embedding = generate_embedding(chunk_data['content'])
            
            # Store embedding in Elasticsearch instead of Supabase
            try:
                index_document_chunk(
                    chunk_id=chunk_record['id'],
                    file_id=file_id,
                    user_id=user_id,
                    content=chunk_data['content'],
                    embedding=embedding,
                    chunk_index=chunk_data['chunk_index'],
                    page_number=chunk_data.get('page_number'),
                    filename=filename
                )
                logger.info(f"✅ Indexed chunk {chunk_record['id']} in Elasticsearch")
            except Exception as es_error:
                logger.error(f"Failed to index chunk in Elasticsearch: {es_error}")
                # Continue processing other chunks even if one fails
        
        return chunk_records
    except Exception as e:
        raise Exception(f"Failed to process file chunks: {str(e)}")

def upload_pdf_file(user_id: str, filename: str, file_content: bytes) -> Dict[str, Any]:
    """Complete file upload process (supports multiple types) with Elasticsearch"""
    try:
        from supabase_client import supabase, get_or_create_user
        if supabase is None:
            raise Exception("Supabase client not initialized")
        
        user_record = get_or_create_user(user_id)
        user_uuid = user_record['id']
        
        extracted_data = extract_text_from_file(file_content, filename)
        content_type = extracted_data.get('mime_type', 'application/octet-stream')
        
        file_path = upload_file_to_storage(file_content, filename, user_uuid)
        file_record = create_file_record(user_uuid, filename, len(file_content), file_path, content_type)
        
        supabase.table('files').update({
            'upload_status': 'processing'
        }).eq('id', file_record['id']).execute()
        
        try:
            # Process chunks with Elasticsearch indexing
            chunk_records = process_file_chunks(
                file_record['id'], 
                extracted_data['chunks'],
                user_uuid,
                filename
            )
            
            supabase.table('files').update({
                'upload_status': 'processed',
                'updated_at': datetime.now().isoformat()
            }).eq('id', file_record['id']).execute()
            
            logger.info(f"✅ File processed: {len(chunk_records)} chunks indexed in Elasticsearch")
            
            return {
                'success': True,
                'file_id': file_record['id'],
                'filename': filename,
                'total_pages': extracted_data.get('page_count'),
                'chunks_created': len(chunk_records),
                'file_path': file_path
            }
        except Exception as processing_error:
            supabase.table('files').update({
                'upload_status': 'failed',
                'processing_error': str(processing_error),
                'updated_at': datetime.now().isoformat()
            }).eq('id', file_record['id']).execute()
            raise processing_error
    except Exception as e:
        logger.error(f"Error uploading file: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_user_files(user_id: str, limit: int = 50) -> List[Dict[str, Any]]:
    """Get files uploaded by a user"""
    try:
        from supabase_client import supabase, get_or_create_user
        
        if supabase is None:
            print("Supabase client not initialized")
            return []
        
        # Get or create user and get their UUID
        user_record = get_or_create_user(user_id)
        user_uuid = user_record['id']
        
        response = supabase.table('files').select('*').eq('user_id', user_uuid).order('created_at', desc=True).limit(limit).execute()
        return response.data if response.data else []
    except Exception as e:
        print(f"Error fetching user files: {e}")
        return []

def get_file_by_id(file_id: str) -> Optional[Dict[str, Any]]:
    """Get file by ID"""
    try:
        from supabase_client import supabase
        
        if supabase is None:
            print("Supabase client not initialized")
            return None
        
        response = supabase.table('files').select('*').eq('id', file_id).execute()
        return response.data[0] if response.data else None
    except Exception as e:
        print(f"Error fetching file: {e}")
        return None

def search_similar_chunks(query: str, user_id: str, limit: int = 5, use_reranking: bool = True) -> List[Dict[str, Any]]:
    """
    Search for similar file chunks using Elasticsearch vector similarity with optional re-ranking
    
    Args:
        query: Search query
        user_id: Firebase user ID
        limit: Number of results to return
        use_reranking: Whether to use cross-encoder re-ranking for better results
        
    Returns:
        List of matching chunks with similarity scores
    """
    try:
        from supabase_client import get_or_create_user
        from elasticsearch_client import search_similar_chunks as es_search
        
        # Map Firebase UID to UUID
        user_record = get_or_create_user(user_id)
        user_uuid = user_record['id']
        
        # Generate query embedding using semantic embeddings
        query_vector = generate_embedding(query)
        
        # Retrieve more candidates for re-ranking (if enabled)
        initial_limit = limit * 3 if use_reranking and SEMANTIC_EMBEDDINGS_AVAILABLE else limit
        
        # Search using Elasticsearch with hybrid search (vector + keyword)
        try:
            results = es_search(
                query_embedding=query_vector,
                user_id=user_uuid,
                k=initial_limit,
                num_candidates=initial_limit * 2,
                use_hybrid=True,
                query_text=query
            )
            
            if results:
                logger.info(f"✅ Elasticsearch returned {len(results)} results")
                
                # Apply re-ranking if enabled and available
                if use_reranking and SEMANTIC_EMBEDDINGS_AVAILABLE and len(results) > 1:
                    try:
                        # Extract documents for re-ranking
                        documents = [r['content'] for r in results]
                        
                        # Re-rank using cross-encoder
                        ranked_indices = rerank_results(query, documents, top_k=limit)
                        
                        # Reorder results based on re-ranking scores
                        reranked_results = []
                        for idx, rerank_score in ranked_indices:
                            result = results[idx].copy()
                            result['rerank_score'] = rerank_score
                            result['original_similarity'] = result['similarity_score']
                            result['similarity_score'] = rerank_score  # Use rerank score as primary
                            reranked_results.append(result)
                        
                        logger.info(f"✅ Re-ranked {len(results)} results to top {len(reranked_results)}")
                        return reranked_results
                        
                    except Exception as rerank_error:
                        logger.warning(f"⚠️  Re-ranking failed, using Elasticsearch scores: {rerank_error}")
                        return results[:limit]
                
                return results[:limit]
            else:
                logger.warning("No results from Elasticsearch")
                return []
                
        except Exception as e:
            logger.error(f"Elasticsearch search failed: {e}")
            return []
        
    except Exception as e:
        logger.error(f"Error searching similar chunks: {e}")
        return []

def delete_file(file_id: str, user_id: str) -> bool:
    """Delete file and all related data from Supabase and Elasticsearch"""
    try:
        from supabase_client import supabase, get_or_create_user
        from elasticsearch_client import delete_file_chunks
        
        if supabase is None:
            logger.error("Supabase client not initialized")
            return False
        
        # Get or create user and get their UUID
        user_record = get_or_create_user(user_id)
        user_uuid = user_record['id']
        
        # Get file record
        file_record = get_file_by_id(file_id)
        if not file_record or file_record['user_id'] != user_uuid:
            return False
        
        # Delete from Elasticsearch first
        try:
            delete_file_chunks(file_id)
            logger.info(f"✅ Deleted chunks from Elasticsearch for file {file_id}")
        except Exception as e:
            logger.warning(f"Could not delete chunks from Elasticsearch: {e}")
        
        # Delete from storage
        try:
            supabase.storage.from_("files").remove([file_record['file_path']])
        except Exception as e:
            logger.warning(f"Could not delete file from storage: {e}")
        
        # Delete file record (cascade will handle chunks in Supabase)
        supabase.table('files').delete().eq('id', file_id).execute()
        
        logger.info(f"✅ Deleted file {file_id} completely")
        return True
        
    except Exception as e:
        logger.error(f"Error deleting file: {e}")
        return False
